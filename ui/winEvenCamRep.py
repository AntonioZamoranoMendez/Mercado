from tkinter import *
from tkinter import ttk
from models.toplevel import TopWindow
from models.datagrid import DataGrid
from database.database import Database
import tempfile
import webbrowser
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import openpyxl
from openpyxl.utils import get_column_letter
import os
from docx import Document


class WinEventCamRep(TopWindow):
    def __init__(self, master):
        super().__init__(master, title="Reporte de Eventos por Cámara", width=800, height=600, not_rezisable=False)

        # Frame superior
        top_frame = ttk.Frame(self, padding=10)
        top_frame.pack(fill="x", pady=10)

        ttk.Label(top_frame, text="Cámara:").pack(side="left", padx=(0, 5))
        self.cmb_camera = ttk.Combobox(top_frame, state="readonly", width=25)
        self.cmb_camera.pack(side="left", padx=(0, 15))

        btn_generate = ttk.Button(top_frame, text="Generar", command=self.on_generate)
        btn_pdf = ttk.Button(top_frame, text="PDF", command=self.on_generate_pdf)
        btn_word = ttk.Button(top_frame, text="WORD", command=self.on_generate_word)
        btn_excel = ttk.Button(top_frame, text="EXCEL", command=self.on_generate_excel)
        btn_exit = ttk.Button(top_frame, text="Salir", command=self.destroy)

        btn_generate.pack(side="left", expand=True, padx=5)
        btn_pdf.pack(side="left", expand=True, padx=5)
        btn_word.pack(side="left", expand=True, padx=5)
        btn_excel.pack(side="left", expand=True, padx=5)
        btn_exit.pack(side="left", expand=True, padx=5)

        # Contenedor principal
        main_frame = ttk.Frame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        dg_frame = ttk.Frame(main_frame)
        dg_frame.pack(fill="both", expand=True)

        # DataGrid
        self.datagrid = DataGrid(dg_frame, width=780, height=500)
        self.datagrid.add_columns(
            columns=["ID", "Cámara ID", "Fecha/Hora", "Descripción"],
            widths=[50, 100, 200, 400]
        )

        # Base de datos
        self.db = Database()

        # Cargar cámaras y eventos
        self.load_cameras()
        self.load_events()

        # Vincular cambio en cámara
        self.cmb_camera.bind("<<ComboboxSelected>>", lambda e: self.load_events(self.get_selected_camera_id()))

    # ---------------------------- FUNCIONES AUXILIARES ----------------------------

    def load_cameras(self):
        """Carga las cámaras disponibles en el combo"""
        cameras = self.db.get_all_cameras()
        self.camera_map = {}
        names = ["Todas las cámaras"]  # opción global

        for cam in cameras:
            self.camera_map[cam.name] = cam.id
            names.append(cam.name)

        self.cmb_camera["values"] = names
        self.cmb_camera.current(0)  # valor por defecto

    def get_selected_camera_id(self):
        """Obtiene el ID de la cámara seleccionada o None si se elige 'Todas'"""
        selected = self.cmb_camera.get()
        if selected == "Todas las cámaras" or not selected:
            return None
        return self.camera_map.get(selected)

    def load_events(self, camera_id=None):
        """Carga los eventos al DataGrid (filtrando si hay cámara seleccionada)"""
        self.datagrid.clear()
        if camera_id:
            events = self.db.get_events_by_camera(camera_id)
        else:
            events = self.db.get_events()

        for ev in events:
            self.datagrid.insert_row([
                str(ev.id),
                str(ev.camera_id),
                ev.timestamp,
                ev.description
            ])

    # ---------------------------- REPORTES ----------------------------

    def on_generate(self):
        """Genera un reporte HTML filtrado por cámara"""
        camera_id = self.get_selected_camera_id()
        if camera_id:
            events = self.db.get_events_by_camera(camera_id)
        else:
            events = self.db.get_events()

        html_content = self._generate_html_table(events)
        with tempfile.NamedTemporaryFile("w", delete=False, suffix=".html", encoding="utf-8") as f:
            f.write(html_content)
            webbrowser.open(f.name)

    def _generate_html_table(self, events):
        """Crea el contenido HTML de la tabla de eventos"""
        html = """
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Reporte de Eventos</title>
            <style>
                table {border-collapse: collapse; width: 90%; margin: 20px auto;}
                th, td {border: 1px solid #000; padding: 8px; text-align: center;}
                th {background-color: #f2f2f2;}
            </style>
        </head>
        <body>
        <h2 style="text-align:center;">Reporte de Eventos por Cámara</h2>
        <table>
            <tr>
                <th>ID</th><th>Cámara ID</th><th>Fecha/Hora</th><th>Descripción</th>
            </tr>
        """
        for ev in events:
            html += f"<tr><td>{ev.id}</td><td>{ev.camera_id}</td><td>{ev.timestamp}</td><td>{ev.description}</td></tr>"
        html += "</table></body></html>"
        return html

    def on_generate_pdf(self):
        """Genera un PDF con los eventos"""
        camera_id = self.get_selected_camera_id()
        if camera_id:
            events = self.db.get_events_by_camera(camera_id)
        else:
            events = self.db.get_events()
        # Crear archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file.close()
        # Registrar fuente Unicode (usa una que exista en Windows)
        try:
            pdfmetrics.registerFont(TTFont("DejaVu", "C:\\Windows\\Fonts\\DejaVuSans.ttf"))
            font_name = "DejaVu"
        except:
            # Si no existe esa fuente, usa Arial (segura en Windows)
            pdfmetrics.registerFont(TTFont("Arial", "C:\\Windows\\Fonts\\arial.ttf"))
            font_name = "Arial"
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(temp_file.name, pagesize=letter)
        width, height = letter
        c.setFont(font_name, 14)
        c.drawCentredString(width / 2, height - 50, "Reporte de Eventos por Cámara")
        c.setFont(font_name, 11)
        y = height - 80
        c.drawString(40, y, "ID")
        c.drawString(80, y, "Cámara")
        c.drawString(160, y, "Fecha/Hora")
        c.drawString(320, y, "Descripción")
        c.line(40, y - 2, 560, y - 2)
        y -= 20
        c.setFont(font_name, 10)
        for ev in events:
            c.drawString(40, y, str(ev.id))
            c.drawString(80, y, str(ev.camera_id))
            c.drawString(160, y, str(ev.timestamp))
            c.drawString(320, y, ev.description[:60])
            y -= 15
            if y < 50:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 50
        c.save()
        # Abrir el PDF con el visor predeterminado del sistema
        try:
            os.startfile(temp_file.name)
        except Exception:
            # En caso de que os.startfile falle (Linux/mac), usar webbrowser
            webbrowser.open(f"file://{temp_file.name}")

    def on_generate_excel(self):
        """Genera un archivo Excel con los eventos"""
        camera_id = self.get_selected_camera_id()
        if camera_id:
            events = self.db.get_events_by_camera(camera_id)
        else:
            events = self.db.get_events()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Reporte de Eventos"

        headers = ["ID", "Cámara ID", "Fecha/Hora", "Descripción"]
        ws.append(headers)

        for ev in events:
            ws.append([ev.id, ev.camera_id, ev.timestamp, ev.description])

        for i, col in enumerate(ws.columns, start=1):
            max_length = max(len(str(cell.value)) if cell.value else 0 for cell in col)
            ws.column_dimensions[get_column_letter(i)].width = max_length + 2

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        wb.save(temp_file.name)
        temp_file.close()
        os.startfile(temp_file.name)

    def on_generate_word(self):
        """Genera un documento Word con los eventos"""
        camera_id = self.get_selected_camera_id()
        if camera_id:
            events = self.db.get_events_by_camera(camera_id)
        else:
            events = self.db.get_events()

        doc = Document()
        doc.add_heading("Reporte de Eventos por Cámara", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        headers = ["ID", "Cámara ID", "Fecha/Hora", "Descripción"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for ev in events:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ev.id)
            row_cells[1].text = str(ev.camera_id)
            row_cells[2].text = str(ev.timestamp)
            row_cells[3].text = ev.description

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()
        os.startfile(temp_file.name)
