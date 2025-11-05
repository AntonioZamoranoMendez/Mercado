from tkinter import *
from tkinter import ttk
from models.toplevel import TopWindow
from models.datagrid import DataGrid
from database.database import Database
import tempfile
import webbrowser
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile
import webbrowser
import openpyxl
from openpyxl.utils import get_column_letter
import os
from docx import Document

class WinCamerasRep(TopWindow):
    def __init__(self, master):
        super().__init__(master, title="Reporte de camaras", width=700, height=600, not_rezisable=False)

        # Botones superiores
        btns = ttk.Frame(self, padding=10)
        btns.pack(fill="x", pady=10)

        btn_generate = ttk.Button(btns, text="Generar", command=self.on_generate)
        btn_pdf = ttk.Button(btns, text="PDF", command=self.on_generate_pdf)
        btn_word = ttk.Button(btns, text="WORD", command=self.on_generate_word)
        btn_excel = ttk.Button(btns, text="EXCEL", command=self.on_generate_excel)
        btn_exit = ttk.Button(btns, text="Salir", command=self.destroy)
        
        btn_generate.pack(side="left", expand=True, padx=5)
        btn_pdf.pack(side="left", expand=True, padx=5)
        btn_word.pack(side="left", expand=True, padx=5)
        btn_excel.pack(side="left", expand=True, padx=5)
        btn_exit.pack(side="left", expand=True, padx=5)

        # Contenedor principal
        main_frame = ttk.Frame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        dg_frame = ttk.Frame(main_frame)
        dg_frame.pack(fill="both", expand=True)

        # DataGrid
        self.datagrid = DataGrid(dg_frame, width=780, height=500)
        self.datagrid.add_columns(
            columns=["ID", "Cámara", "IP", "Puerto"],
            widths=[50, 200, 250, 150]
        )

        # Base de datos
        self.db = Database()

        # Cargar cámaras
        self.load_cameras()

    def load_cameras(self):
        """Carga todas las cámaras al DataGrid"""
        self.datagrid.clear()
        cameras = self.db.get_all_cameras()
        for cam in cameras:
            self.datagrid.insert_row([str(cam.id), cam.name, cam.ip, str(cam.port)])

    def on_generate(self):
        """Genera una tabla HTML con las cámaras y abre en el navegador"""
        cameras = self.db.get_all_cameras()
        html_content = self._generate_html_table(cameras)
        with tempfile.NamedTemporaryFile("w", delete=False, suffix=".html") as f:
            f.write(html_content)
            webbrowser.open(f.name)

    def _generate_html_table(self, cameras):
        """Crea el contenido HTML de la tabla de cámaras"""
        html = """
        <html>
        <head>
            <title>Reporte de Cámaras</title>
            <style>
                table {border-collapse: collapse; width: 80%; margin: 20px auto;}
                th, td {border: 1px solid #000; padding: 8px; text-align: center;}
                th {background-color: #f2f2f2;}
            </style>
        </head>
        <body>
        <h2 style="text-align:center;">Reporte de Cámaras</h2>
        <table>
            <tr>
                <th>ID</th><th>Cámara</th><th>IP</th><th>Puerto</th>
            </tr>
        """
        for cam in cameras:
            html += f"<tr><td>{cam.id}</td><td>{cam.name}</td><td>{cam.ip}</td><td>{cam.port}</td></tr>"
        html += "</table></body></html>"
        return html

    def on_generate_pdf(self):
        """Genera un PDF con la tabla de cámaras y lo abre en el navegador"""
        cameras = self.db.get_all_cameras()
        # Crear archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file.close()  # cerramos para que reportlab pueda escribirlo

        c = canvas.Canvas(temp_file.name, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - 50, "Reporte de Cámaras")

        c.setFont("Helvetica-Bold", 12)
        y = height - 80
        # encabezado
        c.drawString(50, y, "ID")
        c.drawString(100, y, "Cámara")
        c.drawString(300, y, "IP")
        c.drawString(450, y, "Puerto")
        c.line(50, y - 2, 550, y - 2)
        y -= 20

        c.setFont("Helvetica", 12)
        for cam in cameras:
            c.drawString(50, y, str(cam.id))
            c.drawString(100, y, cam.name)
            c.drawString(300, y, cam.ip)
            c.drawString(450, y, str(cam.port))
            y -= 20
            if y < 50:  # nueva página
                c.showPage()
                y = height - 50

        c.save()
        webbrowser.open(f"file://{temp_file.name}")

    def on_generate_excel(self):
        cameras = self.db.get_all_cameras()

        # Crear libro y hoja
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Reporte de Cámaras"

        # Encabezados
        headers = ["ID", "Cámara", "IP", "Puerto"]
        ws.append(headers)

        # Datos
        for cam in cameras:
            ws.append([cam.id, cam.name, cam.ip, cam.port])

        # Ajustar ancho de columnas
        for i, col in enumerate(ws.columns, start=1):
            max_length = max(len(str(cell.value)) for cell in col)
            ws.column_dimensions[get_column_letter(i)].width = max_length + 2

        # Guardar en archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        wb.save(temp_file.name)
        temp_file.close()

        # Abrir Excel automáticamente
        os.startfile(temp_file.name)

    def on_generate_word(self):
        cameras = self.db.get_all_cameras()
    
        # Crear documento
        doc = Document()
        doc.add_heading("Reporte de Cámaras", level=1)
    
        # Agregar tabla
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
    
        # Encabezados
        hdr_cells = table.rows[0].cells
        headers = ["ID", "Cámara", "IP", "Puerto"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
    
        # Datos
        for cam in cameras:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cam.id)
            row_cells[1].text = cam.name
            row_cells[2].text = cam.ip
            row_cells[3].text = str(cam.port)
    
        # Guardar en archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()
    
        # Abrir Word automáticamente
        os.startfile(temp_file.name)
    